import pdfplumber
import camelot
import re
from collections import defaultdict
from src.preprocessing.cleaning import preprocess_text
from docx import Document

class DigitalPDFPipeline:
    def __init__(self, filepath, y_tolerance=3):
        self.filepath = filepath
        self.y_tolerance = y_tolerance
        self.text_by_page = []
        self.cleaned_text = ""
        self.tables = []

    def extract_text(self):
        clustered_lines = []

        with pdfplumber.open(self.filepath) as pdf:
            for page_index, page in enumerate(pdf.pages):
                words = page.extract_words()

                # Deleting the table extraction part
                # 1. Getting the bounding boxes
                table_bboxes = []
                try:
                    table_settings = {"vertical_strategy": "lines", "horizontal_strategy": "lines"}
                    tables = page.find_tables(table_settings)
                    table_bboxes = [table.bbox for table in tables]
                except Exception as e:
                    print(f"Warning: Could not find tables with pdfplumber on page {page_index + 1}: {e}")

                # 2. Filter out words present from the table
                filtered_words = []
                for word in words:
                    x0, y0, x1, y1 = word['x0'], word['top'], word['x1'], word['bottom']
                    in_table = False
                    for bbox in table_bboxes:
                        x_min, y_min, x_max, y_max = bbox
                        if x0 >= x_min and x1 <= x_max and y0 >= y_min and y1 <= y_max:
                            in_table = True
                            break
                    if not in_table:
                        filtered_words.append(word)

                # 3. regular text extraction
                lines = defaultdict(list)
                for word in filtered_words:
                    y = round(word['top'] / self.y_tolerance) * self.y_tolerance
                    lines[y].append((word['x0'], word['text']))

                for y in sorted(lines):
                    line_text = ' '.join(word for _, word in sorted(lines[y], key=lambda x: x[0]))
                    clustered_lines.append(line_text)

        self.text_by_page = clustered_lines
        return self.text_by_page


    def extract_tables(self):
        try:
            tables = camelot.read_pdf(self.filepath, pages='all')
            parsed_tables = []

            for table in tables:
                df = table.df
                if df.shape[0] < 2:
                    continue

                headers = df.iloc[0].tolist()
                table_data = []

                for _, row in df.iloc[1:].iterrows():
                    row_dict = {headers[i]: row[i] for i in range(len(headers))}
                    table_data.append(row_dict)

                parsed_tables.append(table_data)

            self.tables = parsed_tables
            return parsed_tables
        except Exception as e:
            print(f"Table extraction failed: {e}")
            self.tables = []
            return []

    def run_pipeline(self):
        raw_text = self.extract_text()
        cleaned_text = preprocess_text(raw_text)

        tables = self.extract_tables()
        resultant_data = {"text": cleaned_text}

        if tables:
            resultant_data["tables"] = tables

        return resultant_data

class WordDocxPipeline:
    def __init__(self, filepath):
        self.filepath = filepath
        self.raw_text = ""
        self.doc = Document(filepath)

    def extract_text(self):
        paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.raw_text = "\n".join(paragraphs)
        return self.raw_text

    def extract_tables(self):
        extracted_tables = []

        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            extracted_tables.append(rows)

        return extracted_tables

    def run_pipeline(self):
        raw_text = self.extract_text()
        cleaned_text = preprocess_text(raw_text)
        tables = self.extract_tables()
        return {"text": cleaned_text, "tables": tables}
    
class TxtPipeline:
    def __init__(self, filepath):
        self.filepath = filepath
        self.raw_text = ""

    def extract_text(self):
        with open(self.filepath, 'r', encoding='utf-8', errors='ignore') as file:
            self.raw_text = file.read()
        return self.raw_text

    def run_pipeline(self):
        raw_text = self.extract_text()
        cleaned_text = preprocess_text(raw_text)  
        return {"text": cleaned_text}    
