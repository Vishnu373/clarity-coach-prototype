from src.preprocessing.cleaning import preprocess_text
from docx import Document

class WordDocxPipeline:
    def __init__(self, filepath):
        self.filepath = filepath
        self.raw_text = ""
        self.doc = Document(filepath)

    def extract_text(self):
        paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.raw_text = "\n".join(paragraphs)
        return self.raw_text

    def extract_tables(self):
        extracted_tables = []

        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            extracted_tables.append(rows)

        return extracted_tables
        
    def run_pipeline(self):
        raw_text = self.extract_text()
        cleaned_text = preprocess_text(raw_text)
        tables = self.extract_tables()

        return {
            "text": cleaned_text,
            "tables": tables
            }