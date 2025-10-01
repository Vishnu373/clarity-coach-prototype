import logging
from io import BytesIO
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from .preprocessing import preprocess_text

# Configure logging
logger = logging.getLogger(__name__)

class WordPipeline:
    def __init__(self, file_bytes: bytes):
        self.file_bytes = file_bytes

    def extract_text(self):
        """Extract text content from Word document."""
        try:
            doc = Document(BytesIO(self.file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                logger.warning("No text paragraphs found in Word document")
                return ""
                
            logger.info(f"Successfully extracted text from {len(paragraphs)} paragraphs")
            return '\n'.join(paragraphs)
            
        except PackageNotFoundError as e:
            logger.error(f"Invalid Word document format: {e}")
            raise RuntimeError(f"Invalid or corrupted Word document: {e}")
        except Exception as e:
            logger.error(f"Word text extraction failed: {e}")
            raise RuntimeError(f"Failed to extract text from Word document: {e}")

    def extract_tables(self):
        try:
            doc = Document(BytesIO(self.file_bytes))
            tables = []
            
            for table_num, table in enumerate(doc.tables, 1):
                try:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    if table_data:  # Only add non-empty tables
                        tables.append(table_data)
                        logger.debug(f"Extracted table {table_num} with {len(table_data)} rows")
                        
                except Exception as e:
                    logger.warning(f"Failed to extract table {table_num}: {e}")
                    continue
            
            if tables:
                logger.info(f"Successfully extracted {len(tables)} tables from Word document")
            else:
                logger.debug("No tables found in Word document")
                
            return tables
            
        except PackageNotFoundError as e:
            logger.error(f"Invalid Word document format during table extraction: {e}")
            return []
        except Exception as e:
            logger.error(f"Word table extraction failed: {e}")
            return []

    def run_pipeline(self):
        try:
            raw_text = self.extract_text()
            cleaned_text = preprocess_text(raw_text)
            
            result = {"text": cleaned_text}
            
            tables = self.extract_tables()
            if tables:
                result["tables"] = tables
            
            logger.info("Word document extraction pipeline completed successfully")
            return result
            
        except Exception as e:
            logger.error(f"Word document pipeline failed: {e}")
            return {
                "error": f"Word document extraction failed: {str(e)}",
                "text": ""
            }
