from io import BytesIO
from docx import Document
from .preprocessing import preprocess_text

class WordDocxPipeline:
    def __init__(self, file_bytes: bytes):
        self.file_bytes = file_bytes
        self.raw_text = ""
        self.doc = Document(BytesIO(file_bytes))

    def extract_text(self):
        paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.raw_text = "\n".join(paragraphs)
        return self.raw_text

    def extract_tables(self):
        extracted_tables = []
        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            extracted_tables.append(rows)
        return extracted_tables

    def run_pipeline(self):
        raw_text = self.extract_text()
        cleaned_text = preprocess_text(raw_text)
        tables = self.extract_tables()

        result = {"text": cleaned_text}
        if tables:
            result["tables"] = tables

        return result
